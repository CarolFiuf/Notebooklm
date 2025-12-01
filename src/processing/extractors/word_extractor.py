"""
Word Document Extractor using python-docx and docx2txt

Supports:
- .docx files (Office 2007+)
- .doc files (legacy format via docx2txt fallback)
"""
import logging
import re
from typing import Tuple, Dict, Any
from pathlib import Path

from src.utils.exceptions import FileProcessingError
from src.utils.text_utils import clean_text

logger = logging.getLogger(__name__)

class WordExtractor:
    """
    Extract text from Word documents (.docx, .doc)

    Uses python-docx as primary method, docx2txt as fallback
    """

    def __init__(self):
        self.min_text_length = 50
        logger.info("Word Extractor initialized")

    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from Word document

        Args:
            file_path: Path to the Word file

        Returns:
            Tuple[str, Dict[str, Any]]: (extracted_text, metadata)
        """
        try:
            file_path_obj = Path(file_path)
            logger.info(f"Extracting Word document: {file_path_obj.name}")

            # Determine file type
            file_extension = file_path_obj.suffix.lower()

            # Try extraction methods
            extraction_method = None
            docx_error = None

            # Method 1: Try python-docx (best for .docx files)
            if file_extension == '.docx':
                try:
                    text_content, metadata = self._extract_with_python_docx(file_path)
                    extraction_method = 'python-docx'
                except Exception as e:
                    docx_error = e
                    logger.warning(f"python-docx failed: {e}, trying docx2txt...")
                    text_content = None
            else:
                text_content = None
                docx_error = Exception(f"File extension {file_extension} not natively supported by python-docx")

            # Method 2: Fallback to docx2txt (works for both .docx and .doc)
            if text_content is None:
                try:
                    text_content, metadata = self._extract_with_docx2txt(file_path)
                    extraction_method = 'docx2txt'
                except Exception as e2:
                    logger.error(
                        f"Both Word extractors failed for {file_path}:\n"
                        f"  - python-docx error: {docx_error}\n"
                        f"  - docx2txt error: {e2}"
                    )
                    raise FileProcessingError(
                        f"Word extraction failed with all methods. "
                        f"python-docx: {docx_error}, docx2txt: {e2}"
                    )

            # 🔧 SIMPLIFIED: Don't store extraction method or file extension (not needed in Qdrant)

            # Validate extraction
            if len(text_content.strip()) < self.min_text_length:
                logger.warning(f"Very little text extracted from {file_path}")

            # Clean extracted text
            text_content = clean_text(text_content)

            # 🔧 SIMPLIFIED: No extraction stats in metadata (reduces bloat)
            # Just log for debugging
            logger.info(
                f"Extracted {len(text_content)} chars using {extraction_method}"
            )

            return text_content, metadata

        except Exception as e:
            logger.error(f"Error extracting Word document: {e}")
            raise FileProcessingError(f"Word extraction failed: {e}")

    def _extract_with_python_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract using python-docx library (best for .docx)"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for Word document processing. "
                "Install it with: pip install python-docx"
            )

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            # 🔧 NEW: Thêm line break trước paragraph in đậm (thường là tiêu đề Điều)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Check if paragraph is bold (main article title)
                    is_bold = any(run.bold for run in para.runs if run.text.strip())

                    # If bold and starts with "Điều X.", force new line
                    if is_bold and re.match(r'Điều\s+\d+\.', text, re.IGNORECASE):
                        # Ensure previous paragraph ends properly
                        if paragraphs and not paragraphs[-1].endswith('\n'):
                            paragraphs[-1] += '\n'

                    paragraphs.append(text)

            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_texts.append(" | ".join(row_text))

            # Combine all text
            all_text = "\n\n".join(paragraphs)
            if table_texts:
                all_text += "\n\n=== Tables ===\n\n" + "\n".join(table_texts)

            # 🔧 SIMPLIFIED: No document metadata (reduces bloat in Qdrant)
            # Only return empty dict - legal_extractor will add relevant fields
            metadata = {}

            return all_text, metadata

        except Exception as e:
            raise FileProcessingError(f"python-docx extraction failed: {e}")

    def _extract_with_docx2txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract using docx2txt library (fallback, works for .doc and .docx)"""
        try:
            import docx2txt
        except ImportError:
            raise ImportError(
                "docx2txt is required as fallback for Word document processing. "
                "Install it with: pip install docx2txt"
            )

        try:
            # Extract text
            text_content = docx2txt.process(file_path)

            # 🔧 SIMPLIFIED: No metadata (reduces bloat)
            metadata = {}

            return text_content, metadata

        except Exception as e:
            raise FileProcessingError(f"docx2txt extraction failed: {e}")
